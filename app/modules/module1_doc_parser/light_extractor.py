"""
module1_doc_parser/light_extractor.py
----------------------------------------
Step 1 (fallback) — minimal-overhead, dependency-light text extraction.

No upload, no API key, no ML models, no layout/OCR analysis: plain per-format
readers only. Used when the Unstructured hosted API is unavailable, its key
isn't configured, or the caller explicitly asks for the light path (see
document_extractor.extract_document_text).

Supported formats: .pdf, .docx, .xlsx/.xlsm, .csv, .md/.txt (and anything
else text-like, decoded as plain text).
"""

import csv
import io
from pathlib import Path


class LightExtractor:
    """Per-format plain-text extraction with minimal computational overhead."""

    def extract(self, filename: str, content: bytes) -> str:
        """Extract plain text from file bytes, dispatching on file extension."""
        if not content:
            return ""

        suffix = Path(filename).suffix.lower()
        try:
            if suffix == ".pdf":
                return self._extract_pdf(content)
            if suffix == ".docx":
                return self._extract_docx(content)
            if suffix in (".xlsx", ".xlsm"):
                return self._extract_xlsx(content)
            if suffix == ".csv":
                return self._extract_csv(content)
        except Exception:
            return ""

        # .md, .txt, and any other text-like format
        return content.decode("utf-8", errors="replace")

    def _extract_pdf(self, content: bytes) -> str:
        from pypdf import PdfReader

        try:
            reader = PdfReader(io.BytesIO(content))
        except Exception:
            return ""
        parts = []
        for page in reader.pages:
            page_text = (page.extract_text() or "").strip()
            if page_text:
                parts.append(page_text)
        return "\n\n".join(parts)

    def _extract_docx(self, content: bytes) -> str:
        from docx import Document

        document = Document(io.BytesIO(content))
        parts = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n\n".join(parts)

    def _extract_xlsx(self, content: bytes) -> str:
        from openpyxl import load_workbook

        workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        parts = []
        for sheet in workbook.worksheets:
            parts.append(f"# {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(value) for value in row if value is not None]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n\n".join(parts)

    def _extract_csv(self, content: bytes) -> str:
        text = content.decode("utf-8", errors="replace")
        reader = csv.reader(io.StringIO(text))
        return "\n".join(" | ".join(row) for row in reader if any(cell.strip() for cell in row))
