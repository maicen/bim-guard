"""
BIMGUARD AI — Word / PDF Report Generator
modules/report_generator.py

Generates a formatted Word (.docx) compliance report from a completed
BIMGUARD AI compliance run. The report is suitable for client delivery,
design team coordination, or regulatory submission.

Usage in Streamlit:
  from modules.report_generator import generate_word_report
  docx_bytes = generate_word_report(results, impact, project_meta)
  st.download_button("Download Word report", docx_bytes, "bimguard_report.docx")

Dependencies:
  pip install python-docx
"""

import io
from datetime import datetime
from typing import Optional

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ── COLOUR CONSTANTS ──────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1F, 0x38, 0x64) if DOCX_AVAILABLE else None
BLUE   = RGBColor(0x2E, 0x75, 0xB6) if DOCX_AVAILABLE else None
WHITE  = RGBColor(0xFF, 0xFF, 0xFF) if DOCX_AVAILABLE else None
GREY   = RGBColor(0x66, 0x66, 0x66) if DOCX_AVAILABLE else None
RED    = RGBColor(0xA3, 0x2D, 0x2D) if DOCX_AVAILABLE else None
AMBER  = RGBColor(0x85, 0x4F, 0x0B) if DOCX_AVAILABLE else None
GREEN  = RGBColor(0x0F, 0x6E, 0x56) if DOCX_AVAILABLE else None

BAND_COLOURS = {
    "Critical": (0xA3, 0x2D, 0x2D),
    "High":     (0xC2, 0x51, 0x0A),
    "Medium":   (0xB4, 0x53, 0x09),
    "Low":      (0x0F, 0x6E, 0x56),
}


def _set_cell_bg(cell, hex_colour: str):
    """Set table cell background colour using OOXML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour)
    tcPr.append(shd)


def _add_rule(doc, colour: str = "1F3864", size: int = 6):
    """Add a horizontal rule paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), colour)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _heading(doc, text: str, level: int = 1):
    """Add a styled heading paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = NAVY
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = BLUE
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = GREY
    return p


def _body(doc, text: str, italic: bool = False, colour=None):
    """Add a body text paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.italic = italic
    if colour:
        run.font.color.rgb = colour
    return p


def _kv(doc, key: str, value: str):
    """Add a key: value line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    k = p.add_run(key + ": ")
    k.bold = True
    k.font.name = "Arial"
    k.font.size = Pt(10)
    v = p.add_run(str(value))
    v.font.name = "Arial"
    v.font.size = Pt(10)
    return p


def generate_word_report(
    results: list,
    impact,
    project_meta: Optional[dict] = None,
) -> bytes:
    """
    Generate a formatted Word compliance report.

    Parameters
    ----------
    results : list
        List of compliance result dicts from compliance_runner.
    impact : ImpactSummary
        Cost model impact summary from CostModel.calculate_impact().
    project_meta : dict, optional
        Project information: {"project_name", "client", "prepared_by",
                              "ifc_file", "date", "revision"}.

    Returns
    -------
    bytes
        Word document as bytes, ready for st.download_button().
    """
    if not DOCX_AVAILABLE:
        raise ImportError(
            "python-docx is required for report generation. "
            "Install with: pip install python-docx"
        )

    meta = project_meta or {}
    now = datetime.now().strftime("%d %B %Y")

    doc = DocxDocument()

    # ── PAGE SETUP ─────────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ── COVER BLOCK ─────────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(40)
    title.paragraph_format.space_after = Pt(4)
    tr = title.add_run("BIMGUARD AI")
    tr.bold = True
    tr.font.name = "Arial"
    tr.font.size = Pt(28)
    tr.font.color.rgb = NAVY

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(20)
    sr = subtitle.add_run("Corrosion Compliance Report")
    sr.font.name = "Arial"
    sr.font.size = Pt(14)
    sr.font.color.rgb = BLUE
    sr.italic = True

    _add_rule(doc)

    # Project metadata table
    meta_table = doc.add_table(rows=6, cols=2)
    meta_table.style = "Table Grid"
    meta_rows = [
        ("Project",       meta.get("project_name", "—")),
        ("Client",        meta.get("client", "—")),
        ("Prepared by",   meta.get("prepared_by", "BIMGUARD AI v1.0.0")),
        ("IFC file",      meta.get("ifc_file", "—")),
        ("Report date",   meta.get("date", now)),
        ("Revision",      meta.get("revision", "P01")),
    ]
    for i, (k, v) in enumerate(meta_rows):
        row = meta_table.rows[i]
        row.cells[0].text = k
        row.cells[1].text = v
        _set_cell_bg(row.cells[0], "DEEAF1")
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(10)
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # ── EXECUTIVE SUMMARY ───────────────────────────────────────────────────────
    _heading(doc, "1. Executive Summary")
    _add_rule(doc)

    # Count by band
    bands = {"Critical": 0, "High": 0, "Medium": 0, "Low": 0}
    for r in results:
        band = r.get("risk_band", "Low") if isinstance(r, dict) else getattr(r, "risk_band", "Low")
        if band in bands:
            bands[band] += 1

    total = len(results)
    issues = sum(v for k, v in bands.items() if k != "Low")

    _body(doc, f"This report presents the automated corrosion compliance assessment produced by BIMGUARD AI for {total} MEP service elements. The assessment was performed using two corrosion compliance engines: GC-001 (galvanic corrosion, ruleset BIMGUARD-GC-001 v1.0.0) and CC-001 (crevice corrosion, ruleset BIMGUARD-CC-001 v1.0.0).")

    _body(doc, f"The assessment identified {issues} elements at Medium risk or above requiring remediation action. Full details of each flagged element, the composite risk score, the mechanism responsible, and the recommended mitigation are provided in Section 3.")

    # Summary metrics table
    doc.add_paragraph()
    sm = doc.add_table(rows=2, cols=4)
    sm.alignment = WD_TABLE_ALIGNMENT.CENTER
    sm.style = "Table Grid"
    headers = ["Total elements", "Issues flagged", "Est. cost (£)", "Programme delay"]
    values  = [
        str(total),
        str(issues),
        f"£{impact.total_cost_gbp:,.0f}",
        f"{impact.total_days} days",
    ]
    for i, h in enumerate(headers):
        cell = sm.rows[0].cells[i]
        cell.text = h
        _set_cell_bg(cell, "1F3864")
        cell.paragraphs[0].runs[0].font.color.rgb = WHITE
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.name = "Arial"
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, v in enumerate(values):
        cell = sm.rows[1].cells[i]
        cell.text = v
        cell.paragraphs[0].runs[0].font.name = "Arial"
        cell.paragraphs[0].runs[0].font.size = Pt(12)
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Cost avoidance note
    avoidance = impact.total_cost_gbp * 5
    _body(doc,
        f"Based on published industry benchmarks, the cost of resolving these {issues} issues at the current design stage is estimated at £{impact.total_cost_gbp:,.0f}. "
        f"Resolution of equivalent issues discovered during construction would cost approximately 6× more — an estimated £{impact.total_cost_gbp * 6:,.0f} — representing a cost avoidance of £{avoidance:,.0f} by addressing these findings now.",
        colour=GREEN
    )

    # ── RISK DISTRIBUTION ────────────────────────────────────────────────────────
    _heading(doc, "2. Risk Distribution")
    _add_rule(doc)

    rd = doc.add_table(rows=len(bands) + 1, cols=4)
    rd.style = "Table Grid"
    for i, h in enumerate(["Risk band", "Element count", "Estimated cost (£)", "Programme delay"]):
        cell = rd.rows[0].cells[i]
        cell.text = h
        _set_cell_bg(cell, "2E75B6")
        cell.paragraphs[0].runs[0].font.color.rgb = WHITE
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.name = "Arial"
        cell.paragraphs[0].runs[0].font.size = Pt(10)

    for ri, band in enumerate(["Critical", "High", "Medium", "Low"]):
        row = rd.rows[ri + 1]
        bd = impact.issues_by_band.get(band, {"count": bands[band], "cost": 0, "days": 0})
        data = [
            band,
            str(bands[band]),
            f"£{bd.get('cost', 0):,.0f}",
            f"{bd.get('days', 0)} days",
        ]
        col_hex, _ = BAND_COLOURS.get(band, ((0xFF, 0xFF, 0xFF), None))
        hex_str = "".join(f"{c:02X}" for c in BAND_COLOURS.get(band, (0xFF, 0xFF, 0xFF)))

        for ci, val in enumerate(data):
            cell = row.cells[ci]
            cell.text = val
            if ci == 0 and band != "Low":
                _set_cell_bg(cell, hex_str)
                cell.paragraphs[0].runs[0].font.color.rgb = WHITE
                cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.name = "Arial"
            cell.paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # ── ISSUE REGISTER ───────────────────────────────────────────────────────────
    _heading(doc, "3. Issue Register")
    _add_rule(doc)
    _body(doc, "The following table lists all elements flagged at Medium risk or above. Each row contains the IFC GlobalID, service type, material, risk band, composite score, mechanism, and recommended mitigation.")

    doc.add_paragraph()

    # Issue table
    issue_results = [
        r for r in results
        if (r.get("risk_band") if isinstance(r, dict) else getattr(r, "risk_band", "Low")) != "Low"
    ]

    if issue_results:
        cols = ["GlobalID", "Service", "Material", "Band", "Score", "Mechanism", "Mitigation"]
        col_widths = [3.5, 3.0, 3.0, 1.8, 1.5, 1.8, 5.5]  # cm

        ir = doc.add_table(rows=len(issue_results) + 1, cols=len(cols))
        ir.style = "Table Grid"

        for ci, h in enumerate(cols):
            cell = ir.rows[0].cells[ci]
            cell.text = h
            _set_cell_bg(cell, "1F3864")
            run = cell.paragraphs[0].runs[0]
            run.font.color.rgb = WHITE
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(8)

        for ri, r in enumerate(issue_results):
            if isinstance(r, dict):
                gid     = r.get("global_id", r.get("GlobalID", "—"))[:12]
                svc     = r.get("system_type", r.get("element_type", "—"))[:20]
                mat     = r.get("material_label", r.get("material", "—"))[:20]
                band    = r.get("risk_band", "—")
                score   = f"{r.get('composite_score', r.get('GC001_score', 0)):.3f}"
                mech    = r.get("mechanism", r.get("GC001_mechanism", "GC"))
                mit     = r.get("recommended_mitigation", "—")[:60]
            else:
                gid   = str(getattr(r, "global_id", "—"))[:12]
                svc   = str(getattr(r, "system_type", getattr(r, "element_type", "—")))[:20]
                mat   = str(getattr(r, "material_label", getattr(r, "material_key", "—")))[:20]
                band  = str(getattr(r, "risk_band", "—"))
                score = f"{float(getattr(r, 'composite_score', 0)):.3f}"
                mech  = str(getattr(r, "mechanism", "—"))
                mit   = str(getattr(r, "mitigations", ["—"])[0] if getattr(r, "mitigations", []) else "—")[:60]

            vals = [gid, svc, mat, band, score, mech, mit]
            row = ir.rows[ri + 1]
            hex_str = "".join(f"{c:02X}" for c in BAND_COLOURS.get(band, (0xFF, 0xFF, 0xFF)))
            bg = ri % 2 == 0

            for ci, val in enumerate(vals):
                cell = row.cells[ci]
                cell.text = val
                run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(val)
                run.font.name = "Arial"
                run.font.size = Pt(8)
                if ci == 3 and band != "Low":
                    _set_cell_bg(cell, hex_str)
                    run.font.color.rgb = WHITE
                    run.bold = True
                elif bg:
                    _set_cell_bg(cell, "F4F6FB")

    doc.add_paragraph()

    # ── METHODOLOGY ──────────────────────────────────────────────────────────────
    _heading(doc, "4. Assessment Methodology")
    _add_rule(doc)

    _body(doc, "BIMGUARD AI implements a White Box Architecture in which every compliance decision is traceable to a named standard and a specific numerical threshold. The following rulesets and standards were applied in this assessment:")

    methods = [
        ("GC-001 v1.0.0", "Galvanic corrosion — voltage gap vs NASA-STD-6012 thresholds, area ratio risk bands, PREN adequacy check (IMOA Design Manual 4th Ed.)"),
        ("CC-001 v1.0.0", "Crevice corrosion — CCT adequacy per ASTM G48 Method B / CIRIA C692, geometry class per CIBSE Guide G, environment severity per EN ISO 15329:2007"),
        ("IFC parser",    "ifcopenshell — reads IFC 2x3 and IFC4; normalises material names; reads Pset_PipeSegmentOccurrence, Pset_CoveringCommon, Pset_ZoneCommon"),
        ("BCF output",    "BCF 2.1 (buildingSMART) — GlobalID, viewpoint, markup, issue history per element"),
        ("ISO 19650",     "Information container status applied per ISO 19650-2 workflow — S2 / S4 / Status A"),
    ]
    mt = doc.add_table(rows=len(methods) + 1, cols=2)
    mt.style = "Table Grid"
    for i, h in enumerate(["Component", "Standards and approach"]):
        cell = mt.rows[0].cells[i]
        cell.text = h
        _set_cell_bg(cell, "2E75B6")
        run = cell.paragraphs[0].runs[0]
        run.font.color.rgb = WHITE
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10)

    for ri, (comp, desc) in enumerate(methods):
        row = mt.rows[ri + 1]
        for ci, val in enumerate([comp, desc]):
            cell = row.cells[ci]
            cell.text = val
            run = cell.paragraphs[0].runs[0]
            run.font.name = "Arial"
            run.font.size = Pt(10)
            if ci == 0:
                run.bold = True
            if ri % 2 == 0:
                _set_cell_bg(cell, "F4F6FB")

    doc.add_paragraph()

    # ── DISCLAIMER ────────────────────────────────────────────────────────────────
    _heading(doc, "5. Limitations and Disclaimer", level=2)
    _body(doc,
        "This report is produced by an automated compliance checking tool. Results are based on data present in the submitted IFC model and are subject to the accuracy and completeness of that data. Area ratio calculations use nominal diameter estimates where modelled surface areas are not available. BIMGUARD AI does not replace a qualified corrosion engineer's judgement and should be used as a screening tool to prioritise further investigation. All Critical and High findings should be reviewed by a competent engineer before remediation action is taken.",
        italic=True, colour=GREY
    )
    _body(doc,
        f"Ruleset versions: BIMGUARD-GC-001 v1.0.0 | BIMGUARD-CC-001 v1.0.0 | Assessment date: {now}",
        italic=True, colour=GREY
    )

    # ── SAVE TO BYTES ─────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── STREAMLIT INTEGRATION SNIPPET ────────────────────────────────────────────
STREAMLIT_SNIPPET = '''
# Add to Streamlit app — BCF Issue Manager or dedicated Reports page
# ─────────────────────────────────────────────────────────────────────────────
import streamlit as st
from modules.report_generator import generate_word_report
from modules.cost_model import CostModel

st.subheader("Export compliance report")

with st.expander("Project information (optional)", expanded=False):
    col1, col2 = st.columns(2)
    with col1:
        project_name = st.text_input("Project name", value="")
        client       = st.text_input("Client", value="")
        prepared_by  = st.text_input("Prepared by", value="BIMGUARD AI")
    with col2:
        ifc_file  = st.text_input("IFC file reference", value="")
        revision  = st.text_input("Revision", value="P01")

project_meta = {
    "project_name": project_name,
    "client":       client,
    "prepared_by":  prepared_by,
    "ifc_file":     ifc_file,
    "revision":     revision,
}

if st.session_state.get("compliance_results"):
    results = st.session_state.compliance_results
    model   = st.session_state.get("cost_model", CostModel())
    impact  = model.calculate_impact(results)

    if st.button("Generate Word report", type="primary"):
        with st.spinner("Generating report..."):
            try:
                docx_bytes = generate_word_report(results, impact, project_meta)
                st.download_button(
                    label="Download Word report (.docx)",
                    data=docx_bytes,
                    file_name="BIMGUARD_Compliance_Report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
                st.success("Report generated successfully.")
            except ImportError as e:
                st.error(f"python-docx not installed: {e}")
            except Exception as e:
                st.error(f"Report generation failed: {e}")
else:
    st.info("Run a compliance check first to enable report export.")
'''
